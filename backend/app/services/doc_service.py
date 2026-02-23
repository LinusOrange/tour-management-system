import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
import re

TITLE_FONT = '方正小标宋简体'
BODY_FONT = '仿宋'


def _apply_run_font(run, font_name: str, size: Pt | None = None, bold: bool | None = None):
    """统一设置中英文字体，确保 Word 对中文字体生效。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)

    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def _add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    _apply_run_font(run, TITLE_FONT, size=Pt(22) if level == 0 else Pt(16), bold=True)
    return paragraph


def _add_body_paragraph(doc: Document, text: str = ''):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _apply_run_font(run, BODY_FONT, size=Pt(14))
    return paragraph




def _sanitize_ai_section_text(text: str, section_heading: str = '') -> str:
    """清理 AI 输出中的 Markdown 与重复标题。"""
    if not text:
        return ''

    normalized_heading = section_heading.replace('一、', '').replace('二、', '').replace('三、', '').replace('四、', '').replace('五、', '').strip()

    lines = []
    for raw in str(text).splitlines():
        line = raw.strip()
        if not line:
            continue

        # 去除 markdown 前缀 #/##/### 等
        line = re.sub(r'^#{1,6}\s*', '', line)
        # 去除 markdown 列表前缀
        line = re.sub(r'^(?:[-*+]\s+|\d+[\.)]\s+)', '', line)
        # 去除粗体标记
        line = line.replace('**', '')

        # 去除和章节同名的重复标题行
        if normalized_heading and line in {normalized_heading, section_heading}:
            continue

        lines.append(line)

    return '\n'.join(lines).strip()


def _apply_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _apply_run_font(run, BODY_FONT, size=Pt(12))


def create_study_tour_docx(data: dict):
    doc = Document()

    title_text = data.get('title') or '研学课程方案'
    title = _add_heading(doc, f"《{title_text}》", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ai_sections = data.get('ai_sections', {})
    locations = data.get('locations', [])
    timeline = data.get('timeline', [])

    # 第一部分：研学基地（来自排产工作台+AI润色）
    _add_heading(doc, '一、研学基地', level=1)
    section_1 = _sanitize_ai_section_text(ai_sections.get('section_1_base', ''), '一、研学基地')
    if section_1:
        _add_body_paragraph(doc, section_1)
    else:
        for loc in locations:
            _add_body_paragraph(doc, f"基地名称：{loc.get('name', '未命名基地')}")
            _add_body_paragraph(doc, f"基地地址：{loc.get('address', '待补充')}")
            _add_body_paragraph(doc, loc.get('description', '暂无基地描述'))
            _add_body_paragraph(doc, '')

    # 第二/三/四部分：AI 生成
    section_map = [
        ('二、课程背景', 'section_2_background'),
        ('三、研学目标', 'section_3_goals'),
        ('四、课程亮点', 'section_4_highlights')
    ]
    for heading, key in section_map:
        _add_heading(doc, heading, level=1)
        clean_text = _sanitize_ai_section_text(ai_sections.get(key, '内容生成失败，请手动补充。'), heading)
        _add_body_paragraph(doc, clean_text or '内容生成失败，请手动补充。')

    # 第五部分：研学流程（AI流程说明 + 时间轴表格）
    _add_heading(doc, '五、研学流程', level=1)
    process_text = _sanitize_ai_section_text(ai_sections.get('section_5_process', ''), '五、研学流程')
    _add_body_paragraph(doc, process_text)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '时间段'
    hdr_cells[2].text = '活动名称'
    hdr_cells[3].text = '基地/时长'

    for i, act in enumerate(timeline, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = act.get('display_time', '时间待定')
        row_cells[2].text = act.get('title', '未命名活动')
        row_cells[3].text = f"{act.get('location_name', '待定基地')} / {act.get('duration', 0)} 分钟"

    _apply_table_font(table)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
